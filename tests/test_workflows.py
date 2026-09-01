from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document
from openpyxl import Workbook, load_workbook
from pptx import Presentation

from khz_workstation.workflows import document_outline_to_slides, document_table_to_sheet, sheet_range_to_document_table
from khz_workstation.workspace import Workspace


class CrossOfficeWorkflowTests(unittest.TestCase):
    def test_deterministic_cross_office_workflows(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td) / "ws"; ws = Workspace.create(root, "Flow")
            wb = Workbook(); sh = wb.active; sh.title = "Summary"; sh.append(["Name", "Amount"]); sh.append(["A", 10]); sh.append(["B", 20])
            try:
                wb.save(root / "source.xlsx")
            finally:
                wb.close()
            out_doc = sheet_range_to_document_table(ws, "source.xlsx", "Summary", "A1:B3", "range.docx")
            doc = Document(out_doc); self.assertEqual(doc.tables[0].cell(1, 1).text, "10")
            out_sheet = document_table_to_sheet(ws, "range.docx", 0, "table.xlsx")
            check = load_workbook(out_sheet, data_only=False)
            try:
                self.assertEqual(check["ImportedTable"]["B3"].value, "20")
            finally:
                check.close()
            d = Document(); d.add_heading("Quarterly Review", 1); d.add_heading("Operations", 2); d.add_heading("Finding A", 3); d.add_heading("Finance", 2); d.save(root / "outline.docx")
            out_ppt = document_outline_to_slides(ws, "outline.docx", "outline.pptx")
            prs = Presentation(out_ppt); self.assertGreaterEqual(len(prs.slides), 3)
            events = [e["what"] for e in ws.audit.read()]
            self.assertIn("workflow.sheet_range_to_document_table", events)
            self.assertIn("workflow.document_table_to_sheet", events)
            self.assertIn("workflow.document_outline_to_slides", events)


if __name__ == "__main__":
    unittest.main()
